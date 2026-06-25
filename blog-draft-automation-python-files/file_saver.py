"""생성된 초안을 txt 또는 docx 파일로 저장하는 모듈입니다."""

from __future__ import annotations

from pathlib import Path


def save_txt(path: str, content: str) -> str:
    """텍스트 파일로 저장합니다."""
    target = Path(path)
    target.write_text(content, encoding="utf-8")
    return str(target)


def save_docx(path: str, content: str) -> str:
    """python-docx가 설치되어 있으면 docx 파일로 저장합니다."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("docx 저장을 위해 python-docx 설치가 필요합니다.") from exc

    document = Document()
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
        elif stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        else:
            document.add_paragraph(stripped)

    target = Path(path)
    document.save(target)
    return str(target)
